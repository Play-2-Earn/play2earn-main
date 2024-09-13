from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import os
import fitz  # PyMuPDF for PDF processing
from docx import Document
import re
import logging

# Configure logging
logging.basicConfig(level=logging.DEBUG)

def allowed_file(filename, allowed_extensions):
    """Check if the uploaded file is allowed based on its extension."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions

def preprocess_text(text):
    """Preprocess text by converting it to lowercase."""
    return text.lower()

def calculate_context_score(jd, cv):
    """Calculate the context score between a job description and a CV using cosine similarity."""
    jd_processed = preprocess_text(jd)
    cv_processed = preprocess_text(cv)
    documents = [jd_processed, cv_processed]
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform(documents)
    cosine_sim = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])
    return cosine_sim[0][0]

def process_cv_file(file_path, file_name):
    """Process the uploaded CV file and extract text based on the file type."""
    text = ""
    if file_name.lower().endswith('.pdf'):
        try:
            with fitz.open(file_path) as pdf:
                for page in pdf:
                    text += page.get_text()
        except Exception as e:
            logging.error(f"Error processing PDF file {file_name}: {e}")
            return file_name, text, None, None, None, None, None
    elif file_name.lower().endswith('.docx'):
        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + '\n'
        except Exception as e:
            logging.error(f"Error processing .docx file {file_name}: {e}")
            return file_name, text, None, None, None, None, None
    elif file_name.lower().endswith('.doc'):
        try:
            import docx
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + '\n'
        except Exception as e:
            logging.error(f"Error processing .doc file {file_name}: {e}")
            return file_name, text, None, None, None, None, None
    else:
        logging.error(f"Unsupported file type: {file_name}")
        return file_name, text, None, None, None, None, None

    info = extract_information(text)
    if info is None:
        return file_name, text, None, None, None, None, None

    if len(info) != 5:
        return file_name, text, None, None, None, None, None

    name, designation, experience, education, skills = info
    return file_name, text, name, designation, experience, education, skills

def extract_information(text):
    """Extract relevant information such as name, designation, experience, education, and skills from the CV text."""
    name = extract_name(text)
    designation = extract_designation(text)
    experience = extract_experience(text)
    education = extract_education(text)
    skills = extract_skills(text)
    
    return name, designation, experience, education, skills

def extract_name(text):
    """Extract the name of the candidate from the CV text."""
    lines = text.splitlines()
    for line in lines[:10]:  # Check the first 10 lines for the name
        if line.strip() and len(line.strip().split()) <= 3:  # Assume name is short, e.g., "John Doe"
            return line.strip()
    return None

def extract_designation(text):
    """Extract the designation or job title from the CV text."""
    patterns = ['senior', 'manager', 'developer', 'engineer', 'specialist', 'director', 'consultant']
    for pattern in patterns:
        match = re.search(r'\b{}\b'.format(pattern), text, re.IGNORECASE)
        if match:
            return match.group()
    return None

def extract_experience(text):
    """Extract the experience details from the CV text."""
    match = re.search(r'\b\d+\+?\s*(years?|months?)\s*(of)?\s*(experience|work)', text, re.IGNORECASE)
    if match:
        return match.group()
    return None

def extract_education(text):
    """Extract the education details from the CV text."""
    patterns = ['Bachelor', 'Master', 'PhD', 'BSc', 'MSc', 'MBA', 'Diploma']
    for pattern in patterns:
        match = re.search(r'\b{}\b'.format(pattern), text, re.IGNORECASE)
        if match:
            return match.group()
    return None

def extract_skills(text):
    """Extract the skills from the CV text."""
    skill_keywords = ['Python', 'Java', 'C++', 'Project Management', 'Data Analysis', 'SQL', 'Marketing']
    found_skills = []
    for keyword in skill_keywords:
        # Use re.escape() to handle special characters in skill names
        if re.search(r'\b{}\b'.format(re.escape(keyword)), text, re.IGNORECASE):
            found_skills.append(keyword)
    return ', '.join(found_skills) if found_skills else None

def calculate_skill_score(cv_skills, job_description):
    """Calculate skill match score based on the job description and CV skills."""
    job_description = preprocess_text(job_description)
    cv_skills = preprocess_text(cv_skills)
    return calculate_context_score(job_description, cv_skills)

def calculate_experience_score(cv_experience, job_experience):
    """Calculate experience match score based on the job description and CV experience."""
    cv_experience = preprocess_text(cv_experience)
    job_experience = preprocess_text(job_experience)
    return calculate_context_score(job_experience, cv_experience)
